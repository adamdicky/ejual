from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docx" / "product-and-catalogue-management-report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_margins(section) -> None:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def apply_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size in [("Title", 22), ("Heading 1", 14), ("Heading 2", 11.5)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        if style_name != "Title":
            style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def add_code_block(document: Document, code: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.6)
    set_cell_shading(cell, "F3F4F6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(code.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    document.add_paragraph()


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].width = Inches(widths[idx])
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(hdr[idx], "E5E7EB")
        p = hdr[idx].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(value)
            run.font.size = Pt(9)

    document.add_paragraph()


def bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)


doc = Document()
set_margins(doc.sections[0])
apply_base_styles(doc)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Product and Catalogue Management Component Report")
run.bold = True
run.font.size = Pt(22)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(18)
subtitle.add_run("eJual | Seller Module under /seller/products").italic = True

for line in [
    "Selected component: seller Product and Catalogue Management",
    "Coverage: product list, create, edit, variants, product images, and inventory",
    "Prepared as coursework evidence for AI-assisted development, debugging, refactoring, and documentation",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.add_run(line)

doc.add_section(WD_SECTION.NEW_PAGE)
set_margins(doc.sections[-1])

doc.add_heading("1. Component Overview", level=1)
doc.add_paragraph(
    "The selected component is the seller-facing Product and Catalogue Management module in eJual. "
    "Its purpose is to let an authenticated seller create a clothing product, define purchasable variants, "
    "attach product-level images, and maintain inventory records for each variant. This component was selected "
    "because it is the first complete vertical slice of the eJual marketplace and it demonstrates both UI work "
    "with Mantine and server-side data handling with Payload CMS."
)
doc.add_paragraph(
    "The main routes are `/seller/products`, `/seller/products/new`, and `/seller/products/[id]/edit`. "
    "The most important implementation files are `src/app/(frontend)/seller/products/actions.ts`, "
    "`src/app/(frontend)/seller/products/_components/ProductForm.tsx`, "
    "`src/app/(frontend)/seller/products/_components/ProductEditWorkspace.tsx`, and the route files that load "
    "seller-owned products. The module depends on the existing Payload collections `products`, `categories`, "
    "`product-variants`, `product-images`, `inventory`, `media`, and `users`."
)
bullet(doc, "Mantine-first seller dashboard UI with consistent form and tab patterns.")
bullet(doc, "Payload Local API for reads and server actions for mutations.")
bullet(doc, "User-scoped operations with `overrideAccess: false` and authenticated `user` context.")
bullet(doc, "Seller ownership checks before editing products, images, variants, or inventory.")
bullet(doc, "`reservedQuantity` kept system-managed rather than seller-editable.")

doc.add_heading("2. Code for Selected Component", level=1)
doc.add_paragraph(
    "The following excerpts represent the core flow of the component. They show how the module validates seller "
    "input, creates seller-owned products, protects ownership-sensitive updates, and structures the editor "
    "workspace into manageable tabs."
)

doc.add_paragraph(
    "Excerpt 1 is from `actions.ts` and shows the server action entry points plus the shared helper that limits "
    "product access to the owning seller. This is central to the module because every create or edit request goes "
    "through these actions."
)
add_code_block(
    doc,
    """
const requireOwnedProduct = async (payload, user, productID) => {
  const { docs } = await payload.find({
    collection: 'products',
    limit: 1,
    overrideAccess: false,
    user,
    where: {
      and: [{ id: { equals: productID } }, { seller: { equals: user.id } }],
    },
  })

  if (!docs.length) throw new Error('Product not found.')
  return docs[0]
}

export const updateProduct = async (productID, _prevState, formData) => {
  const user = await requireActiveSeller()
  const payload = await getPayload({ config: configPromise })
  await requireOwnedProduct(payload, user, productID)
  ...
}
""",
)

doc.add_paragraph(
    "Excerpt 2 is the product creation flow. It creates the `products` record first, then loops through every "
    "submitted variant, creates a `product-variants` record, and immediately creates the matching `inventory` "
    "record. This keeps the product, its purchasable options, and its stock records in sync."
)
add_code_block(
    doc,
    """
const colors = formData.getAll('color')
if (!colors.length) throw new Error('At least one variant is required.')

const product = await payload.create({
  collection: 'products',
  data: {
    basePrice: numberFromForm(formData, 'basePrice'),
    category: numberFromForm(formData, 'category'),
    description: stringFromForm(formData, 'description'),
    featured: formData.get('featured') === 'on',
    productName: stringFromForm(formData, 'productName'),
    seller: user.id,
    status: 'available',
  },
  overrideAccess: false,
  user,
})

for (let index = 0; index < colors.length; index++) {
  const variant = await payload.create({
    collection: 'product-variants',
    data: {
      additionalPrice: numberAt(formData, 'additionalPrice', index),
      color: stringAt(formData, 'color', index),
      product: product.id,
      size: stringAt(formData, 'size', index),
    },
    overrideAccess: false,
    user,
  })

  await payload.create({
    collection: 'inventory',
    data: {
      reservedQuantity: 0,
      stockQuantity: numberAt(formData, 'stockQuantity', index),
      variant: variant.id,
    },
    overrideAccess: false,
    user,
  })
}
""",
)

doc.add_paragraph(
    "Excerpt 3 is from `ProductForm.tsx`. It shows how the create page supports multiple variants and product-level "
    "image uploads. This excerpt matters because the seller workflow would be incomplete if the seller could only "
    "create one variant or had to leave the create flow to add core catalogue details."
)
add_code_block(
    doc,
    """
const [variantCount, setVariantCount] = useState(1)

<Button onClick={() => setVariantCount((count) => count + 1)} type="button" variant="light">
  Add another variant
</Button>

{Array.from({ length: variantCount }, (_, index) => (
  <Box key={index}>
    <Text>{index === 0 ? 'First variant' : `Variant ${index + 1}`}</Text>
    <TextInput label="Color" name="color" required />
    <TextInput label="Size" name="size" required />
    <NumberInput label="Additional price" name="additionalPrice" prefix="RM " required />
    <NumberInput label="Stock quantity" name="stockQuantity" required />
  </Box>
))}

<FileInput
  label="Upload product images"
  name="productImages"
  multiple
  accept="image/*"
/>
""",
)

doc.add_paragraph(
    "Excerpt 4 is from `ProductEditWorkspace.tsx`. It separates product maintenance into tabs for details, variants, "
    "images, and inventory. The tab structure reduces cognitive load by showing only one maintenance surface at a "
    "time while still keeping the whole product configuration in one page."
)
add_code_block(
    doc,
    """
<Tabs defaultValue="details">
  <Tabs.List>
    <Tabs.Tab value="details">Details</Tabs.Tab>
    <Tabs.Tab value="variants">Variants</Tabs.Tab>
    <Tabs.Tab value="images">Images</Tabs.Tab>
    <Tabs.Tab value="inventory">Inventory</Tabs.Tab>
  </Tabs.List>

  <Tabs.Panel value="details"><ProductForm ... /></Tabs.Panel>
  <Tabs.Panel value="variants"><AddVariantForm ... /></Tabs.Panel>
  <Tabs.Panel value="images"><ProductImageOrderForm ... /></Tabs.Panel>
  <Tabs.Panel value="inventory"><InventoryForm ... /></Tabs.Panel>
</Tabs>
""",
)

doc.add_heading("3. AI Prompt Logs", level=1)
doc.add_paragraph(
    "Instead of pasting a raw transcript, this section presents a curated prompt log grouped by development purpose. "
    "This better demonstrates how AI was used as a coding assistant and where human review was still necessary."
)
add_table(
    doc,
    ["Prompt", "Purpose", "AI output summary", "Student correction or refinement"],
    [
        [
            "Build a seller-facing Product and Catalogue Management interface using Mantine UI with routes for list, create, and edit.",
            "Code generation and feature scaffolding",
            "Proposed the seller routes, Mantine-first form layout, table/list surfaces, and Payload Local API usage.",
            "The generated direction was accepted, but the student constrained it to clothing-only categories, removed Payload chrome from app routes, and required theme consistency.",
        ],
        [
            "At seller/products/new, there should be a button of add another variant under First variant.",
            "UI improvement",
            "Expanded the create form from a single variant block to a repeatable variant entry pattern.",
            "The student reviewed the form labels, kept Mantine components only, and made the first product status implicit rather than editable during creation.",
        ],
        [
            "I tried to create a product but got basePrice must be a valid number. I typed 20 and it automated into RM20.00.",
            "Debugging",
            "Identified that the formatted Mantine number input was submitting a string with the `RM` prefix and that the action layer had to normalize it.",
            "The final code strips `RM` and commas before numeric conversion and applies the same rule to repeated variant price fields.",
        ],
        [
            "A user can only see their own products.",
            "Security and ownership correction",
            "Recommended ownership checks for seller read and write paths.",
            "The final module enforces seller scoping on the edit route and reuses `requireOwnedProduct` in mutating actions.",
        ],
        [
            "Create a .READme file for eJual's github, for other collaborators to read, and set it up on their machine.",
            "Documentation generation",
            "Produced a collaborator-oriented setup guide with routes, collections, and development rules.",
            "The student kept the content aligned to the actual repo and current auth/UI rules.",
        ],
    ],
    [1.8, 1.0, 1.5, 1.8],
)

doc.add_paragraph(
    "Overall, AI was most useful for fast first drafts of route structure, form composition, and candidate fixes. "
    "It was only partially useful for domain constraints, because the student still had to correct categories, seller "
    "ownership expectations, and exact UX behavior for the eJual clothing marketplace."
)

doc.add_heading("4. AI-Assisted Debugging", level=1)
doc.add_paragraph(
    "Two real debugging cases from this component are documented below. In both cases, AI helped explain the likely "
    "cause, but the final fix still required checking the actual form submission and access flow used by this repo."
)

doc.add_heading("Issue 1: `basePrice must be a valid number` during product creation", level=2)
doc.add_paragraph(
    "Problem. The create product form used Mantine `NumberInput` with an `RM ` prefix for Malaysian Ringgit. "
    "When the seller entered a value such as `20`, the submitted payload could contain a formatted string such as "
    "`RM 20.00`. A direct numeric conversion failed and the server action returned the error `basePrice must be a valid number.`"
)
doc.add_paragraph(
    "How AI helped. AI correctly pointed to the mismatch between formatted UI values and server-side parsing, which "
    "narrowed the debugging effort to the action helpers instead of the database schema."
)
doc.add_paragraph("Before fix:")
add_code_block(
    doc,
    """
const numberFromForm = (formData: FormData, key: string): number => {
  const value = Number(formData.get(key))
  if (!Number.isFinite(value)) throw new Error(`${key} must be a valid number.`)
  return value
}
""",
)
doc.add_paragraph("After fix:")
add_code_block(
    doc,
    """
const numberFromForm = (formData: FormData, key: string): number => {
  const rawValue = formData.get(key)
  const normalizedValue = String(rawValue || '')
    .trim()
    .replace(/^RM\\s*/i, '')
    .replace(/,/g, '')
  const value = Number(normalizedValue)
  if (!Number.isFinite(value)) throw new Error(`${key} must be a valid number.`)
  return value
}
""",
)
doc.add_paragraph(
    "Why the correction works. The fix normalizes the formatted string before conversion, so the server action can "
    "accept the display format produced by the UI while still validating that the final value is numeric."
)

doc.add_heading("Issue 2: seller inventory updates needed explicit ownership validation", level=2)
doc.add_paragraph(
    "Problem. Earlier `updateInventory` logic updated an inventory row by ID after only checking that the user was "
    "logged in. That was a logic and security weakness because a seller action should only touch inventory records "
    "for variants that belong to that seller's own product."
)
doc.add_paragraph(
    "How AI helped. AI suggested centralizing seller ownership checks. The student then verified the actual data model "
    "and implemented the check against `products`, `product-variants`, and `inventory` rather than trusting the incoming IDs."
)
doc.add_paragraph("Before fix:")
add_code_block(
    doc,
    """
export const updateInventory = async (productID, inventoryID, _prevState, formData) => {
  const user = await requireActiveSeller()
  const payload = await getPayload({ config: configPromise })

  await payload.update({
    id: inventoryID,
    collection: 'inventory',
    data: {
      stockQuantity: numberFromForm(formData, 'stockQuantity'),
    },
    overrideAccess: false,
    user,
  })
}
""",
)
doc.add_paragraph("After fix:")
add_code_block(
    doc,
    """
await requireOwnedProduct(payload, user, productID)
const { docs: variants } = await payload.find({
  collection: 'product-variants',
  limit: 500,
  overrideAccess: false,
  user,
  where: { product: { equals: productID } },
})
const variantIDs = variants.map((variant) => variant.id)

const { docs: inventoryItems } = await payload.find({
  collection: 'inventory',
  limit: 1,
  overrideAccess: false,
  user,
  where: {
    and: [{ id: { equals: inventoryID } }, { variant: { in: variantIDs } }],
  },
})

if (!inventoryItems.length) throw new Error('Inventory record not found.')
""",
)
doc.add_paragraph(
    "Why the correction works. The updated action proves that the target inventory row belongs to one of the current "
    "seller's variants before performing the update. This prevents cross-product or cross-seller modification through "
    "manually constructed requests."
)

doc.add_heading("5. Code Smells and Refactoring", level=1)
doc.add_paragraph(
    "This component also contained maintainability issues that were not immediate runtime failures. Two refactoring "
    "cases are discussed below."
)

doc.add_heading("Refactoring 1: duplicated variant creation flow", level=2)
doc.add_paragraph(
    "Smell. The module contains repeated logic for creating a product variant and its matching inventory row. In the "
    "current code, `createProduct` and `addVariant` both create a variant first and then create an `inventory` record "
    "with `reservedQuantity: 0`. This duplication increases the chance of future inconsistency if one path changes "
    "without the other."
)
doc.add_paragraph(
    "Refactoring technique. Extract method. A small helper such as `createVariantWithInventory(...)` would reduce "
    "duplication while preserving the existing behavior and Payload API usage."
)
doc.add_paragraph("Representative repeated code:")
add_code_block(
    doc,
    """
const variant = await payload.create({
  collection: 'product-variants',
  data: {
    additionalPrice: ...,
    color: ...,
    product: productID,
    size: ...,
  },
  overrideAccess: false,
  user,
})

await payload.create({
  collection: 'inventory',
  data: {
    reservedQuantity: 0,
    stockQuantity: ...,
    variant: variant.id,
  },
  overrideAccess: false,
  user,
})
""",
)
doc.add_paragraph(
    "Improvement. Extracting the shared sequence would make variant creation easier to test and would keep inventory "
    "defaults consistent across both create and edit flows."
)

doc.add_heading("Refactoring 2: rigid single-variant create form evolved into a flexible form", level=2)
doc.add_paragraph(
    "Smell. The earlier create page supported only one initial variant. That design was too rigid for clothing items, "
    "where a seller typically needs several size or color combinations at creation time. The result was unnecessary "
    "extra work and an incomplete first-step workflow."
)
doc.add_paragraph(
    "Refactoring technique. Replace fixed structure with controlled repetition. The create form was changed from one "
    "hard-coded variant block into a state-driven repeated Mantine form section."
)
doc.add_paragraph("Before:")
add_code_block(
    doc,
    """
<Stack gap=\"sm\">
  <Text fw={700}>First variant</Text>
  <Group grow>
    <TextInput label=\"Color\" name=\"color\" required />
    <TextInput label=\"Size\" name=\"size\" required />
  </Group>
  <Group grow>
    <NumberInput label=\"Additional price\" name=\"additionalPrice\" prefix=\"RM \" required />
    <NumberInput label=\"Stock quantity\" name=\"stockQuantity\" required />
  </Group>
</Stack>
""",
)
doc.add_paragraph("After:")
add_code_block(
    doc,
    """
const [variantCount, setVariantCount] = useState(1)

<Button onClick={() => setVariantCount((count) => count + 1)} type=\"button\">
  Add another variant
</Button>

{Array.from({ length: variantCount }, (_, index) => (
  <Box key={index}>
    <Text>{index === 0 ? 'First variant' : `Variant ${index + 1}`}</Text>
    ...
  </Box>
))}
""",
)
doc.add_paragraph(
    "Improvement. The new form better matches the eJual domain, where one clothing product commonly has multiple "
    "variants. It also reduces the need to immediately jump into the edit page for basic setup."
)

doc.add_heading("6. Component Documentation", level=1)
doc.add_paragraph(
    "Purpose. The Product and Catalogue Management component is the seller workspace for maintaining clothing listings "
    "in eJual. It covers the full catalogue lifecycle from initial product creation to later edits of images, variants, "
    "and stock quantities."
)
doc.add_paragraph(
    "Routes. `/seller/products` lists only the signed-in seller's products. `/seller/products/new` creates a new product "
    "with at least one variant and optional product images. `/seller/products/[id]/edit` updates product details and "
    "exposes separate tabs for variants, images, and inventory."
)
doc.add_paragraph(
    "Main files. `actions.ts` contains all server mutations, seller validation, image upload attachment, and route "
    "revalidation. `ProductForm.tsx` is the shared create and edit form for core product data. "
    "`ProductEditWorkspace.tsx` composes the edit-page tabs and the image ordering workflow. Route files load seller-"
    "scoped data with the Payload Local API."
)
doc.add_paragraph(
    "Data dependencies. The component writes to `products`, `product-variants`, `inventory`, and `product-images`, and "
    "uploads files into `media`. It reads `categories` for seller selection and `users` for the authenticated seller "
    "context. All seller-scoped mutations use `overrideAccess: false` with the authenticated `user` object."
)
doc.add_paragraph(
    "Usage constraints. Sellers must be authenticated and have `accountStatus === 'active'` to mutate data. Sellers may "
    "only view and edit their own products. `reservedQuantity` is system-managed for cart and checkout reservation logic "
    "and must not be edited manually by the seller. Product images are attached at the product level, not the variant level."
)
doc.add_paragraph(
    "UI constraints. Mantine UI is the primary component system. Seller pages should avoid Payload template chrome, keep "
    "one dominant action per surface, and remain responsive on both desktop and mobile. Product edit flows should continue "
    "to use tabs rather than stacking every management function into one long page."
)

doc.add_heading("7. Conclusion", level=1)
doc.add_paragraph(
    "For this component, AI was useful for generating first-pass code structure, explaining likely causes of runtime "
    "errors, and suggesting improvements to the seller workflow. However, the AI output still required correction in "
    "three areas: eJual domain rules, access control expectations, and the exact user experience details needed for a "
    "clothing marketplace. The final implementation quality came from combining AI assistance with direct inspection of "
    "the repo, targeted fixes, and student review of the generated output."
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
